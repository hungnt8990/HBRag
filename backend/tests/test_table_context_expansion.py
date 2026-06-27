from __future__ import annotations

import asyncio
from io import BytesIO
from types import SimpleNamespace
from uuid import UUID, uuid4

from docx import Document

from app.services.parsers.parser_docx_parser import DocxParser
from app.services.rag.rag_answer_service import ContextChunk, RagAnswerService
from app.services.chunkers.chunker_table_aware_chunking import table_aware_chunk_text

DOCUMENT_ID = UUID("10000000-aaaa-4000-9000-000000000001")
OTHER_DOCUMENT_ID = UUID("10000000-aaaa-4000-9000-000000000002")
TABLE_QUERY = (
    "Danh sách hộ gia đình được kêu gọi ủng hộ khắc phục thiệt hại do cơn bão số 10 "
    "gồm những ai?"
)
HOUSEHOLD_NAMES = [
    "Đoàn Tấn Châu",
    "Khương Thanh Hà",
    "Trần Đình Thanh",
    "Đặng Thị Hiền",
    "Nguyễn Thị Thu Lan",
    "Đoàn Ngọc Bích Thủy",
    "Huỳnh Thị Thanh",
    "Lê Thị Đan Thanh",
    "Nguyễn Văn Tường",
    "Trần Thị Thu Trang",
    "Nguyễn Tấn Trung",
    "Kiều Phước Sen",
    "Nguyễn Thị Hương",
    "Hoàng Quang",
    "Trần Thị Lệ",
    "Trần Nhật Lợi",
    "Lê Văn Hà",
    "Trần Lê Nhật Trâm",
    "Đinh Thị Thúy Phương",
    "Trịnh Minh Quang",
    "Huỳnh Văn Thu",
    "Nguyễn Thị Giang Phương",
    "Lê Thị Thanh",
    "Ngô Anh Đức",
    "Đỗ Thị Thu Hiệp",
    "Nguyễn Văn Liên",
    "Nguyễn Thị Hùng",
    "Phạm Viết Trang",
    "Trần Khanh",
    "Lê Thị Hồng Minh",
    "Trần Thị Đà",
    "Phan Thị Xuân",
    "Trần Thị Hoa",
    "Cao Thị Thuỳ Trang",
]


def _build_household_docx() -> bytes:
    document = Document()
    document.add_paragraph(
        "DANH SÁCH HỘ GIA ĐÌNH ỦNG HỘ KHẮC PHỤC THIỆT HẠI DO CƠN BÃO SỐ 10 "
        "GÂY RA. THƯ KÊU GỌI MTTP NGÀY 6/10/2025"
    )
    document.add_paragraph("TỔ 40/HTT - PHƯỜNG HÒA CƯỜNG")
    table = document.add_table(rows=1, cols=5)
    headers = ["TT", "Họ và tên", "Địa chỉ", "Số tiền(đ)", "Ghi chú"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header

    amounts = ["600.000đ"] * 34
    for index, (name, amount) in enumerate(zip(HOUSEHOLD_NAMES, amounts, strict=True), start=1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = name
        row[2].text = f"Địa chỉ {index}"
        row[3].text = amount
        row[4].text = ""

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "Tổng cộng"
    total_row[2].text = ""
    total_row[3].text = "20.400.000đ"
    total_row[4].text = ""

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _chunk_namespace(chunk: dict, *, document_id: UUID = DOCUMENT_ID):
    return SimpleNamespace(
        id=uuid4(),
        document_id=document_id,
        chunk_index=chunk["chunk_index"],
        content=chunk["content"],
        chunk_metadata=chunk["metadata"],
    )


def _fixture_chunks():
    parsed = DocxParser().parse(_build_household_docx()).text
    raw_chunks, _ = table_aware_chunk_text(parsed, chunk_size=1200)
    chunks = [_chunk_namespace(chunk) for chunk in raw_chunks]
    return parsed, raw_chunks, chunks


def test_docx_parser_serializes_household_table_rows() -> None:
    parsed = DocxParser().parse(_build_household_docx()).text

    assert "TABLE_TITLE table_id=docx_t1" in parsed
    assert parsed.count("TABLE_ROW table_id=docx_t1") >= 35
    assert all(name in parsed for name in HOUSEHOLD_NAMES)
    assert "20.400.000đ" in parsed


def test_table_aware_chunking_emits_required_table_metadata() -> None:
    _parsed, raw_chunks, _chunks = _fixture_chunks()
    row_chunks = [
        chunk for chunk in raw_chunks if chunk["metadata"].get("chunk_type") == "table_row"
    ]
    block_chunks = [
        chunk for chunk in raw_chunks if chunk["metadata"].get("chunk_type") == "table_block"
    ]

    assert len(row_chunks) >= 35
    assert block_chunks
    first_row = row_chunks[0]
    metadata = first_row["metadata"]
    assert metadata["chunk_type"] == "table_row"
    assert metadata["chunk_mode"] == "table_aware"
    assert metadata["table_id"] == "docx_t1"
    assert metadata["table_title"]
    assert metadata["headers"] == ["TT", "Họ và tên", "Địa chỉ", "Số tiền(đ)", "Ghi chú"]
    assert metadata["row_index"] == 1
    assert metadata["row_start"] == 1
    assert metadata["row_end"] == 1


def test_table_expansion_fetches_all_rows_when_retrieval_only_hits_header() -> None:
    _parsed, _raw_chunks, chunks = _fixture_chunks()
    primary = next(
        chunk for chunk in chunks if chunk.chunk_metadata["chunk_type"] == "table_header"
    )
    table_neighbors = [chunk for chunk in chunks if chunk.id != primary.id]

    class FakeRepoWithTable:
        def __init__(self) -> None:
            self.requested_document_id = None

        async def get_table_chunks(self, *, document_id, table_id, exclude_ids):
            self.requested_document_id = document_id
            assert table_id == "docx_t1"
            return [
                chunk
                for chunk in table_neighbors
                if chunk.document_id == document_id and chunk.id not in set(exclude_ids)
            ]

    repo = FakeRepoWithTable()
    service = RagAnswerService(
        chat_repository=repo,  # type: ignore[arg-type]
        reranking_service=SimpleNamespace(),  # type: ignore[arg-type]
        llm_provider=SimpleNamespace(),  # type: ignore[arg-type]
    )

    expanded = asyncio.run(
        service._expand_with_neighbors(
            query=TABLE_QUERY,
            context_chunks=[ContextChunk(citation_index=1, chunk=primary)],
            max_context_chars=3000,
        )
    )

    contents = "\n".join(chunk.chunk.content for chunk in expanded)
    assert repo.requested_document_id == DOCUMENT_ID
    assert all(name in contents for name in HOUSEHOLD_NAMES)
    assert "20.400.000đ" in contents
    assert sum("TABLE_ROW table_id=docx_t1" in chunk.chunk.content for chunk in expanded) >= 35


def test_table_list_query_prompt_contains_all_rows_total_and_no_other_document() -> None:
    _parsed, _raw_chunks, chunks = _fixture_chunks()
    primary = next(chunk for chunk in chunks if chunk.chunk_metadata["chunk_type"] == "table_title")
    other_doc_chunk = SimpleNamespace(
        id=uuid4(),
        document_id=OTHER_DOCUMENT_ID,
        chunk_index=1,
        content="TABLE_ROW table_id=docx_t1 row=1 | Họ và tên: PoC ThinkLabs",
        chunk_metadata={"chunk_type": "table_row", "table_id": "docx_t1"},
    )

    class FakeRepoWithScopedTable:
        async def get_table_chunks(self, *, document_id, table_id, exclude_ids):
            assert document_id == DOCUMENT_ID
            assert document_id != OTHER_DOCUMENT_ID
            return [chunk for chunk in chunks if chunk.id != primary.id] + [other_doc_chunk]

    service = RagAnswerService(
        chat_repository=FakeRepoWithScopedTable(),  # type: ignore[arg-type]
        reranking_service=SimpleNamespace(),  # type: ignore[arg-type]
        llm_provider=SimpleNamespace(),  # type: ignore[arg-type]
    )
    expanded = asyncio.run(
        service._expand_with_neighbors(
            query=TABLE_QUERY,
            context_chunks=[ContextChunk(citation_index=1, chunk=primary)],
            max_context_chars=3000,
        )
    )
    same_document_expanded = [
        chunk for chunk in expanded if chunk.chunk.document_id == DOCUMENT_ID
    ]
    prompt = RagAnswerService._build_user_prompt(
        query=TABLE_QUERY,
        context_chunks=same_document_expanded,
    )

    assert "ENTITY_MATCHED_ROWS:" in prompt
    assert all(name in prompt for name in HOUSEHOLD_NAMES)
    assert "20.400.000đ" in prompt
    assert "PoC ThinkLabs" not in prompt
    assert prompt.count("Đoàn Tấn Châu") == 1


def test_table_row_deduplication_keeps_distinct_rows() -> None:
    _parsed, _raw_chunks, chunks = _fixture_chunks()
    row_chunks = [chunk for chunk in chunks if chunk.chunk_metadata["chunk_type"] == "table_row"]
    duplicated = [
        ContextChunk(citation_index=1, chunk=row_chunks[0]),
        ContextChunk(citation_index=2, chunk=row_chunks[0]),
        ContextChunk(citation_index=3, chunk=row_chunks[1]),
        ContextChunk(citation_index=4, chunk=row_chunks[2]),
    ]
    service = RagAnswerService(
        chat_repository=SimpleNamespace(),  # type: ignore[arg-type]
        reranking_service=SimpleNamespace(),  # type: ignore[arg-type]
        llm_provider=SimpleNamespace(),  # type: ignore[arg-type]
    )

    deduplicated = service._deduplicate_context_chunks(duplicated)

    assert len(deduplicated) == 3
    assert [item.citation_index for item in deduplicated] == [1, 2, 3]
    assert "Đoàn Tấn Châu" in deduplicated[0].chunk.content
    assert "Khương Thanh Hà" in deduplicated[1].chunk.content
    assert "Trần Đình Thanh" in deduplicated[2].chunk.content
