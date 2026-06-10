import sys
from datetime import datetime
from types import SimpleNamespace
from uuid import UUID

from fastapi.testclient import TestClient

from app.api.routes.documents import get_document_repository, get_storage_client
from app.main import app
from app.services.document_parser_service import DocumentParserService
from app.services.parsers import (
    DocumentParser,
    MarkdownParser,
    ParsedDocument,
    ParsedElement,
    PdfParser,
    TextParser,
)
from app.services.parsers.pdf_parser import PdfParser as PdfParserImpl

DOCUMENT_ID = UUID("22222222-2222-2222-2222-222222222222")


class FakeDocumentRepository:
    def __init__(
        self,
        *,
        filename: str = "sample.txt",
        mime_type: str = "text/plain",
        storage_path: str = "documents/sample.txt",
        document_status: str = "uploaded",
    ) -> None:
        self.document = SimpleNamespace(
            id=DOCUMENT_ID,
            status=document_status,
            document_metadata={},
        )
        self.document_file = SimpleNamespace(
            filename=filename,
            mime_type=mime_type,
            storage_path=storage_path,
        )
        self.committed = False
        self.rolled_back = False

    async def get_document(self, document_id: UUID) -> SimpleNamespace | None:
        if document_id != DOCUMENT_ID:
            return None
        return self.document

    async def get_primary_document_file(self, document_id: UUID) -> SimpleNamespace | None:
        if document_id != DOCUMENT_ID:
            return None
        return self.document_file

    async def update_document_status(
        self,
        document: SimpleNamespace,
        status: str,
    ) -> SimpleNamespace:
        document.status = status
        return document

    async def update_document_parsed_content(
        self,
        document: SimpleNamespace,
        *,
        parsed_text: str,
        parsed_at: datetime,
        status: str = "parsed",
    ) -> SimpleNamespace:
        document.parsed_text = parsed_text
        document.parsed_at = parsed_at
        document.status = status
        return document

    async def update_document_metadata(
        self,
        document: SimpleNamespace,
        metadata: dict,
    ) -> SimpleNamespace:
        document.document_metadata = {**document.document_metadata, **metadata}
        return document

    async def commit(self) -> None:
        self.committed = True

    async def rollback(self) -> None:
        self.rolled_back = True


class FakeStorageClient:
    def __init__(self, files: dict[str, bytes]) -> None:
        self.files = files
        self.downloads: list[str] = []

    async def put_file(self, **_: object) -> str:
        raise AssertionError("Upload should not be called by parse tests.")

    async def get_file(self, *, object_name: str) -> bytes:
        self.downloads.append(object_name)
        return self.files[object_name]

    async def delete_file(self, **_: object) -> None:
        raise AssertionError("Delete should not be called by parse tests.")


def test_parse_txt_document() -> None:
    repository = FakeDocumentRepository(
        filename="sample.txt",
        mime_type="text/plain",
        storage_path="documents/sample.txt",
    )
    storage = FakeStorageClient({"documents/sample.txt": b"Hello parser\nSecond line"})
    app.dependency_overrides[get_document_repository] = lambda: repository
    app.dependency_overrides[get_storage_client] = lambda: storage

    try:
        client = TestClient(app)
        response = client.post(f"/api/documents/{DOCUMENT_ID}/parse")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    payload = response.json()
    assert payload["document_id"] == str(DOCUMENT_ID)
    assert payload["status"] == "parsed"
    assert payload["character_count"] == len("Hello parser\nSecond line")
    assert payload["preview"] == "Hello parser\nSecond line"
    assert repository.document.status == "parsed"
    assert repository.document.parsed_text == "Hello parser\nSecond line"
    assert isinstance(repository.document.parsed_at, datetime)
    assert repository.committed is True
    assert repository.rolled_back is False
    assert storage.downloads == ["documents/sample.txt"]


def test_parse_removes_nul_characters_before_storing() -> None:
    repository = FakeDocumentRepository(
        filename="nul.txt",
        mime_type="text/plain",
        storage_path="documents/nul.txt",
    )
    storage = FakeStorageClient({"documents/nul.txt": b"\x00Hello\x00 parser"})
    app.dependency_overrides[get_document_repository] = lambda: repository
    app.dependency_overrides[get_storage_client] = lambda: storage

    try:
        client = TestClient(app)
        response = client.post(f"/api/documents/{DOCUMENT_ID}/parse")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    payload = response.json()
    assert payload["character_count"] == len("Hello parser")
    assert payload["preview"] == "Hello parser"
    assert repository.document.parsed_text == "Hello parser"
    assert repository.committed is True
    assert repository.rolled_back is False


def test_parse_md_document() -> None:
    repository = FakeDocumentRepository(
        filename="notes.md",
        mime_type="text/markdown",
        storage_path="documents/notes.md",
    )
    storage = FakeStorageClient({"documents/notes.md": b"# Heading\n\nSome notes"})
    app.dependency_overrides[get_document_repository] = lambda: repository
    app.dependency_overrides[get_storage_client] = lambda: storage

    try:
        client = TestClient(app)
        response = client.post(f"/api/documents/{DOCUMENT_ID}/parse")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "parsed"
    assert payload["character_count"] == len("# Heading\n\nSome notes")
    assert payload["preview"] == "# Heading\n\nSome notes"
    assert repository.committed is True


def test_parse_rejects_unsupported_document_file() -> None:
    repository = FakeDocumentRepository(
        filename="spreadsheet.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        storage_path="documents/spreadsheet.xlsx",
    )
    storage = FakeStorageClient({"documents/spreadsheet.xlsx": b"not parsed"})
    app.dependency_overrides[get_document_repository] = lambda: repository
    app.dependency_overrides[get_storage_client] = lambda: storage

    try:
        client = TestClient(app)
        response = client.post(f"/api/documents/{DOCUMENT_ID}/parse")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 415
    assert repository.document.status == "uploaded"
    assert repository.committed is False
    assert repository.rolled_back is False
    assert storage.downloads == []


def test_parser_selection_uses_extension_or_mime_type() -> None:
    service = DocumentParserService(
        repository=FakeDocumentRepository(),
        storage=FakeStorageClient({}),
    )

    assert isinstance(
        service.select_parser(filename="sample.txt", mime_type="application/octet-stream"),
        TextParser,
    )
    assert isinstance(
        service.select_parser(filename="README", mime_type="text/markdown"),
        MarkdownParser,
    )
    assert isinstance(
        service.select_parser(filename="report.pdf", mime_type=None),
        PdfParser,
    )

def test_parse_document_stores_structured_parse_metadata() -> None:
    class StructuredParser(DocumentParser):
        supported_extensions = frozenset({".structured"})

        def parse(self, file_content: bytes) -> ParsedDocument:
            return ParsedDocument(
                text=file_content.decode("utf-8"),
                metadata={"source": "fixture"},
                elements=[
                    ParsedElement(
                        element_type="heading",
                        text="Section A",
                        page_number=1,
                        heading_path=["Section A"],
                    )
                ],
            )

    repository = FakeDocumentRepository(
        filename="sample.structured",
        mime_type="application/octet-stream",
        storage_path="documents/sample.structured",
    )
    service = DocumentParserService(
        repository=repository,
        storage=FakeStorageClient({"documents/sample.structured": b"Section A\nBody"}),
        parsers=(StructuredParser(),),
    )

    import anyio

    response = anyio.run(service.parse_document, DOCUMENT_ID)

    assert response.status == "parsed"
    assert repository.document.document_metadata["parser"] == "structured"
    assert repository.document.document_metadata["parsed_metadata"] == {"source": "fixture"}
    assert repository.document.document_metadata["parsed_elements"][0]["element_type"] == "heading"


def _build_docx_with_tables(target_chars: int) -> bytes:
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("CHUONG I QUY DINH CHUNG")
    doc.add_paragraph("Dieu 1. Pham vi dieu chinh")
    doc.add_paragraph("Van ban nay quy dinh noi dung mo dau.")

    table = doc.add_table(rows=0, cols=2)
    cell_text = "Noi dung dong bang dai hon de dat so ky tu can thiet. " * 4
    while True:
        row = table.add_row().cells
        row[0].text = "Tieu chi"
        row[1].text = cell_text
        accumulated = sum(len(paragraph.text) for paragraph in doc.paragraphs)
        for current_table in doc.tables:
            for current_row in current_table.rows:
                for cell in current_row.cells:
                    accumulated += len(cell.text)
        if accumulated >= target_chars:
            break

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_parse_docx_extracts_tables_and_does_not_truncate() -> None:
    docx_bytes = _build_docx_with_tables(target_chars=4500)
    repository = FakeDocumentRepository(
        filename="legal.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        storage_path="documents/legal.docx",
    )
    storage = FakeStorageClient({"documents/legal.docx": docx_bytes})
    app.dependency_overrides[get_document_repository] = lambda: repository
    app.dependency_overrides[get_storage_client] = lambda: storage

    try:
        client = TestClient(app)
        response = client.post(f"/api/documents/{DOCUMENT_ID}/parse")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    payload = response.json()
    assert payload["character_count"] > 4000
    assert len(payload["preview"]) <= 500
    assert repository.document.parsed_text is not None
    assert len(repository.document.parsed_text) == payload["character_count"]
    assert "TABLE_ROW table_id=docx_t1" in repository.document.parsed_text
    assert "Tieu chi" in repository.document.parsed_text


def test_docx_parser_keeps_empty_cells_with_generic_headers() -> None:
    from io import BytesIO

    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Alice"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "Platform"
    table.cell(1, 0).text = "Bob"
    table.cell(1, 1).text = "QA"
    table.cell(1, 2).text = ""

    buffer = BytesIO()
    doc.save(buffer)

    repository = FakeDocumentRepository(
        filename="matrix.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        storage_path="documents/matrix.docx",
    )
    storage = FakeStorageClient({"documents/matrix.docx": buffer.getvalue()})
    app.dependency_overrides[get_document_repository] = lambda: repository
    app.dependency_overrides[get_storage_client] = lambda: storage

    try:
        client = TestClient(app)
        response = client.post(f"/api/documents/{DOCUMENT_ID}/parse")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert "cell_1: Alice" in repository.document.parsed_text
    assert "cell_2:" in repository.document.parsed_text
    assert "cell_3: Platform" in repository.document.parsed_text


def test_parse_pdf_serializes_detected_tables(monkeypatch) -> None:
    class FakePage:
        def extract_text(self, extraction_mode=None):
            return (
                "Name    Area\n"
                "Nguyen Quang Lam    Infrastructure\n"
                "Tran Van An    QA\n"
            )

    class FakeReader:
        def __init__(self, _stream):
            self.pages = [FakePage()]

    monkeypatch.setattr("app.services.parsers.pdf_parser.PdfReader", FakeReader)

    parsed = PdfParserImpl().parse(b"%PDF-1.4 fake")

    assert "TABLE_ROW table_id=pdf_p1_1 page=1 row=1" in parsed.text
    assert "Name: Nguyen Quang Lam" in parsed.text
    assert "Area: Infrastructure" in parsed.text


def test_table_serialization_preserves_multiline_cells_and_width() -> None:
    from app.services.parsers.table_serialization import serialize_table

    rows = [
        ["STT", "Nhom nhiem vu", "Don vi", "Danh sach"],
        [
            "3",
            "Xay dung nen tang RAG tren du lieu noi bo",
            "PTUD",
            "1. Tong Phuoc Lam\n2. Nguyen Quang Lam\n3. Nguyen Trong Hung",
        ],
        [
            "4",
            "Xay dung dich vu OCR dung chung",
            "PM",
            "1. Trinh Thanh Tinh\n2. Duong Sinh Sinh\n3. Nguyen Quang Lam",
        ],
        [
            "5",
            "Kho du lieu AI dung chung",
            "VH",
            "1. Doan Gia Hy\n2. Vo Van Phuc\n3. Vo Van Hoa\n4. Nguyen Quang Lam",
        ],
        [
            "6",
            "Platform AI",
            "PTUD",
            "Cac nhan su trong ke hoach PoC ThinkLabs:\n"
            "1. Phan Anh Tuan\n2. Tran Huy\n3. Nguyen Quang Lam",
        ],
    ]

    serialized = serialize_table(table_id="fixture_t1", rows=rows, page_number=5)
    table_rows = [
        line for line in serialized.splitlines()
        if line.startswith("TABLE_ROW")
    ]

    expected_header = (
        "TABLE_HEADER table_id=fixture_t1 page=5 | "
        "STT | Nhom nhiem vu | Don vi | Danh sach"
    )
    assert expected_header in serialized
    assert len(table_rows) == 4
    assert all("STT:" in row for row in table_rows)
    assert all("Nhom nhiem vu:" in row for row in table_rows)
    assert all("Don vi:" in row for row in table_rows)
    assert all("Danh sach:" in row for row in table_rows)
    assert any("Platform AI" in row and "Nguyen Quang Lam" in row for row in table_rows)
    assert not any(
        "cell_1: Xay dung nen tang RAG tren du lieu noi bo | cell_2: 2. Nguyen Quang Lam"
        in row
        for row in table_rows
    )





def test_parse_pdf_merges_wrapped_aligned_rows() -> None:
    from app.services.parsers.table_serialization import rewrite_text_with_serialized_tables

    raw_text = (
        "No              Work stream                  Owner             People\n"
        "                                                team\n"
        "        Build RAG platform on internal                  2. Nguyen Quang Lam\n"
        " 3      data                                  PTUD      3. Nguyen Trong Hung\n"
        "                                                           4. Vo Van Hoa\n"
        "        Build shared OCR service                         1. Trinh Thanh Tinh\n"
        " 4      for all teams                          PM        2. Duong Sinh Sinh\n"
        "                                                           3. Nguyen Quang Lam\n"
        " 5      Shared AI data warehouse              VH        8. Nguyen Quang Lam\n"
        " 6      Platform AI                           PTUD      7. Nguyen Quang Lam\n"
    )

    serialized = rewrite_text_with_serialized_tables(
        text=raw_text,
        page_number=5,
        table_id_prefix="pdf_p5",
    )
    rows = [line for line in serialized.splitlines() if line.startswith("TABLE_ROW")]

    assert len(rows) >= 4
    lam_rows = [line for line in rows if "Nguyen Quang Lam" in line]
    assert len(lam_rows) == 4
    assert any(
        "No: 3" in line
        and "Work stream: Build RAG platform on internal data" in line
        and "Owner team: PTUD" in line
        for line in lam_rows
    )
    assert any(
        "No: 4" in line
        and "Work stream: Build shared OCR service for all teams" in line
        and "Owner team: PM" in line
        for line in lam_rows
    )
    assert any(
        "No: 5" in line
        and "Work stream: Shared AI data warehouse" in line
        and "Owner team: VH" in line
        for line in lam_rows
    )
    assert any(
        "No: 6" in line
        and "Work stream: Platform AI" in line
        and "Owner team: PTUD" in line
        for line in lam_rows
    )

def test_table_serialization_row_record_preserves_extra_and_missing_columns() -> None:
    from app.services.parsers.table_serialization import build_table_row_record

    extra_value_record = build_table_row_record(
        table_id="fixture_t2",
        row_index=1,
        headers=["A", "B"],
        values=["one", "two", "three"],
    )
    missing_value_record = build_table_row_record(
        table_id="fixture_t2",
        row_index=2,
        headers=["A", "B", "C"],
        values=["one"],
    )

    assert "A: one" in extra_value_record
    assert "B: two" in extra_value_record
    assert "cell_3: three" in extra_value_record
    assert "A: one" in missing_value_record
    assert "B:" in missing_value_record
    assert "C:" in missing_value_record


def test_pdf_parser_uses_pdfplumber_table_cells(monkeypatch) -> None:
    rows = [
        ["STT", "Nhom nhiem vu", "Don vi", "Danh sach"],
        [
            "3",
            "Xay dung nen tang RAG tren du lieu noi bo",
            "PTUD",
            "1. Tong Phuoc Lam\n2. Nguyen Quang Lam\n3. Nguyen Trong Hung",
        ],
        [
            "4",
            "Xay dung dich vu OCR dung chung",
            "PM",
            "1. Trinh Thanh Tinh\n2. Duong Sinh Sinh\n3. Nguyen Quang Lam",
        ],
        [
            "5",
            "Kho du lieu AI dung chung",
            "VH",
            "1. Doan Gia Hy\n2. Vo Van Phuc\n3. Vo Van Hoa\n4. Nguyen Quang Lam",
        ],
        [
            "6",
            "Platform AI",
            "PTUD",
            "Cac nhan su trong ke hoach PoC ThinkLabs:\n"
            "1. Phan Anh Tuan\n2. Tran Huy\n3. Nguyen Quang Lam",
        ],
    ]

    class FakePage:
        def extract_text(self):
            return "Intro text outside the table."

        def extract_tables(self, *, table_settings):
            if table_settings["vertical_strategy"] == "lines":
                return [
                    [
                        ["Xay dung nen tang RAG tren du"],
                        ["2. Nguyen Quang Lam"],
                    ]
                ]
            return [rows]

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    fake_pdfplumber = SimpleNamespace(open=lambda _stream: FakePdf())
    monkeypatch.setitem(sys.modules, "pdfplumber", fake_pdfplumber)
    monkeypatch.setattr(
        PdfParserImpl,
        "_parse_with_pypdf",
        lambda *_args: (_ for _ in ()).throw(AssertionError("pypdf fallback not expected")),
    )

    parsed = PdfParserImpl().parse(b"%PDF-1.4 fake")

    assert "Intro text outside the table." in parsed.text
    assert parsed.text.count("TABLE_ROW table_id=pdf_p1_1 page=1 row=") == 4
    assert "Xay dung nen tang RAG tren du lieu noi bo" in parsed.text
    assert "Xay dung dich vu OCR dung chung" in parsed.text
    assert "Kho du lieu AI dung chung" in parsed.text
    assert "Platform AI" in parsed.text
    assert all(
        "Nguyen Quang Lam" in line
        for line in parsed.text.splitlines()
        if line.startswith("TABLE_ROW")
    )
    assert "cell_1: Xay dung nen tang RAG tren du | cell_2: 2. Nguyen Quang Lam" not in parsed.text


def test_pdf_parser_cleans_presentation_overlay_lines(monkeypatch) -> None:
    class FakePage:
        def extract_text(self):
            return (
                "\x00Quá trình triển khai\n"
                "QQuuáá ttrrììnnhh ttrriiểểnn kkhhaaii\n"
                "Công cụ chuyển đổi dữ liệu GIS 110kV\n"
            )

        def extract_tables(self, *, table_settings):
            return [[
                [
                    "Nhóm, lớp dữ liệu Sổ tay bao gồm nhóm lớp",
                    "",
                    "01",
                ]
            ]]

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    fake_pdfplumber = SimpleNamespace(open=lambda _stream: FakePdf())
    monkeypatch.setitem(sys.modules, "pdfplumber", fake_pdfplumber)
    monkeypatch.setattr(
        PdfParserImpl,
        "_parse_with_pypdf",
        lambda *_args: (_ for _ in ()).throw(AssertionError("pypdf fallback not expected")),
    )

    parsed = PdfParserImpl().parse(b"%PDF-1.4 fake")

    assert "\x00" not in parsed.text
    assert "QQuuáá" not in parsed.text
    assert parsed.text.count("Quá trình triển khai") == 1
    assert "Công cụ chuyển đổi dữ liệu GIS 110kV" in parsed.text
    assert "TABLE_ROW" not in parsed.text

def test_pdf_parser_prefers_page_fallback_for_orphan_diacritics() -> None:
    primary = "3. DEMO\n\u1ea3 \u01a1\nXin C m n"
    fallback = "3. DEMO\nXin C\u1ea3m \u01a1n"

    assert PdfParserImpl._choose_better_page_text(primary, fallback) == fallback

def test_pdf_parser_keeps_primary_when_fallback_has_overlay_duplicates() -> None:
    primary = "\u1ee8NG D\u1ee4NG TRONG EVNCPC\nC\u00e1c l\u1edbp d\u1eef li\u1ec7u"
    fallback = (
        "\u1ee8\u1ee8\u1ee8NG DNG DNG D\u1ee4\u1ee4\u1ee4NG TRONG EVNCPC"
        "NG TRONG EVNCPCNG TRONG EVNCPC"
    )

    assert PdfParserImpl._choose_better_page_text(primary, fallback) == primary

def test_pdf_parser_prefers_coherent_tables_over_fragmented_text_tables() -> None:
    coherent_table = [
        ["STT", "Mang cong nghe", "Phong", "Nhan su"],
        ["3", "Xay dung nen tang RAG tren du lieu noi bo", "PTUD", "Nguyen Quang Lam"],
        ["4", "Xay dung dich vu OCR dung chung", "PM", "Nguyen Quang Lam"],
    ]
    fragmented_table = [
        ["S", "T", "T", "", "M", "a", "n", "g", "", "N", "h", "a", "n"],
        ["3", "", "", "", "R", "A", "G", "", "", "N", "g", "u", "y"],
        ["", "", "", "", "d", "u", "", "l", "i", "e", "u", "", ""],
    ]

    assert PdfParserImpl._score_extracted_table(coherent_table) > (
        PdfParserImpl._score_extracted_table(fragmented_table)
    )
